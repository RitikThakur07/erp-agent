from typing import List, Dict, Optional
from pathlib import Path
import PyPDF2
from docx import Document
import openpyxl
import pandas as pd


class DocumentParser:
    """Parse various document formats for RAG ingestion."""
    
    @staticmethod
    def parse_pdf(file_path: str) -> Dict[str, any]:
        """Extract text from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    text_content.append({
                        "page": page_num + 1,
                        "content": text
                    })
                
                return {
                    "type": "pdf",
                    "filename": Path(file_path).name,
                    "num_pages": len(pdf_reader.pages),
                    "pages": text_content,
                    "full_text": "\n\n".join([p["content"] for p in text_content])
                }
        except Exception as e:
            return {
                "type": "pdf",
                "filename": Path(file_path).name,
                "error": str(e)
            }
    
    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, any]:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            tables_data = []
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            return {
                "type": "docx",
                "filename": Path(file_path).name,
                "paragraphs": paragraphs,
                "tables": tables_data,
                "full_text": "\n".join(paragraphs)
            }
        except Exception as e:
            return {
                "type": "docx",
                "filename": Path(file_path).name,
                "error": str(e)
            }
    
    @staticmethod
    def parse_xlsx(file_path: str) -> Dict[str, any]:
        """Extract data from Excel file."""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            sheets_data = {}
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Read as DataFrame for easier processing
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                sheets_data[sheet_name] = {
                    "rows": len(df),
                    "columns": list(df.columns),
                    "data": df.to_dict(orient='records'),
                    "text_summary": df.to_string()
                }
            
            return {
                "type": "xlsx",
                "filename": Path(file_path).name,
                "sheets": sheets_data,
                "sheet_names": list(sheets_data.keys())
            }
        except Exception as e:
            return {
                "type": "xlsx",
                "filename": Path(file_path).name,
                "error": str(e)
            }
    
    @staticmethod
    def parse_txt(file_path: str) -> Dict[str, any]:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return {
                "type": "txt",
                "filename": Path(file_path).name,
                "full_text": content
            }
        except Exception as e:
            return {
                "type": "txt",
                "filename": Path(file_path).name,
                "error": str(e)
            }
    
    @classmethod
    def parse_document(cls, file_path: str) -> Dict[str, any]:
        """Auto-detect and parse document based on extension."""
        ext = Path(file_path).suffix.lower()
        
        parsers = {
            '.pdf': cls.parse_pdf,
            '.docx': cls.parse_docx,
            '.xlsx': cls.parse_xlsx,
            '.xls': cls.parse_xlsx,
            '.txt': cls.parse_txt,
            '.md': cls.parse_txt
        }
        
        parser = parsers.get(ext)
        if not parser:
            return {
                "type": "unknown",
                "filename": Path(file_path).name,
                "error": f"Unsupported file type: {ext}"
            }
        
        return parser(file_path)


class DocumentChunker:
    """Split documents into chunks for better retrieval."""
    
    @staticmethod
    def chunk_text(
        text: str, 
        chunk_size: int = 1000, 
        overlap: int = 200
    ) -> List[str]:
        """Split text into overlapping chunks."""
        if not text:
            return []
        
        chunks = []
        start = 0
        text_len = len(text)
        
        while start < text_len:
            end = start + chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence or paragraph boundary
            if end < text_len:
                last_period = chunk.rfind('.')
                last_newline = chunk.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > chunk_size * 0.5:  # At least 50% of chunk
                    chunk = chunk[:break_point + 1]
                    end = start + break_point + 1
            
            chunks.append(chunk.strip())
            start = end - overlap
        
        return [c for c in chunks if c]  # Remove empty chunks
    
    @staticmethod
    def chunk_document(parsed_doc: Dict, chunk_size: int = 1000) -> List[Dict]:
        """Chunk a parsed document into smaller pieces with metadata."""
        chunks = []
        
        if "error" in parsed_doc:
            return chunks
        
        doc_type = parsed_doc["type"]
        filename = parsed_doc["filename"]
        
        if doc_type in ["pdf", "docx", "txt"]:
            full_text = parsed_doc.get("full_text", "")
            text_chunks = DocumentChunker.chunk_text(full_text, chunk_size)
            
            for i, chunk in enumerate(text_chunks):
                chunks.append({
                    "text": chunk,
                    "metadata": {
                        "source": filename,
                        "type": doc_type,
                        "chunk_index": i,
                        "total_chunks": len(text_chunks)
                    }
                })
        
        elif doc_type == "xlsx":
            # For Excel, chunk each sheet
            sheets = parsed_doc.get("sheets", {})
            for sheet_name, sheet_data in sheets.items():
                text = sheet_data.get("text_summary", "")
                text_chunks = DocumentChunker.chunk_text(text, chunk_size)
                
                for i, chunk in enumerate(text_chunks):
                    chunks.append({
                        "text": chunk,
                        "metadata": {
                            "source": f"{filename} - {sheet_name}",
                            "type": doc_type,
                            "sheet": sheet_name,
                            "chunk_index": i,
                            "total_chunks": len(text_chunks)
                        }
                    })
        
        return chunks
